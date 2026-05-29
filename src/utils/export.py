"""
Conversation Export Utilities
==============================

Exports a conversation (messages + citations) to PDF or DOCX.

PDF requires ``reportlab`` (optional dep).  If not installed, the PDF
export raises ``ImportError`` with a helpful message.
DOCX uses ``python-docx`` which is a required dep.
"""
from __future__ import annotations

import io
from typing import Any

from ..utils.logging_config import get_logger

logger = get_logger(__name__)


def export_conversation_docx(
    title: str,
    messages: list[dict[str, Any]],
) -> bytes:
    """Return a DOCX file as bytes for the given conversation messages.

    Each message dict must have keys: role, content, timestamp (optional).
    """
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()
    doc.add_heading(title, level=1)

    for msg in messages:
        role = msg.get('role', 'unknown').capitalize()
        content = msg.get('content', '')
        timestamp = msg.get('timestamp', '')

        para = doc.add_paragraph()
        run_role = para.add_run(f"{role}: ")
        run_role.bold = True
        if role.lower() == 'assistant':
            run_role.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
        run_role.font.size = Pt(11)

        run_content = para.add_run(content)
        run_content.font.size = Pt(11)

        if timestamp:
            ts_para = doc.add_paragraph(timestamp)
            ts_para.runs[0].font.size = Pt(8)
            ts_para.runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

        doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def export_conversation_pdf(
    title: str,
    messages: list[dict[str, Any]],
) -> bytes:
    """Return a PDF file as bytes for the given conversation messages.

    Requires ``reportlab``.  Raises ImportError if not installed.
    """
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import cm
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
    except ImportError as exc:
        raise ImportError(
            "reportlab is required for PDF export. Install it with: pip install reportlab>=4.0"
        ) from exc

    buf = io.BytesIO()
    doc_template = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=2 * cm,
        rightMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
    )

    styles = getSampleStyleSheet()
    title_style = styles['Title']
    user_style = ParagraphStyle(
        'User',
        parent=styles['Normal'],
        textColor=colors.HexColor('#111827'),
        spaceAfter=4,
        leading=16,
    )
    assistant_style = ParagraphStyle(
        'Assistant',
        parent=styles['Normal'],
        textColor=colors.HexColor('#1A56DB'),
        spaceAfter=4,
        leading=16,
    )
    ts_style = ParagraphStyle(
        'Timestamp',
        parent=styles['Normal'],
        textColor=colors.HexColor('#6B7280'),
        fontSize=8,
        spaceAfter=12,
    )

    story = [Paragraph(title, title_style), Spacer(1, 0.5 * cm)]

    for msg in messages:
        role = msg.get('role', 'unknown').capitalize()
        content = msg.get('content', '').replace('\n', '<br/>')
        timestamp = msg.get('timestamp', '')

        msg_style = assistant_style if role.lower() == 'assistant' else user_style
        story.append(Paragraph(f"<b>{role}:</b> {content}", msg_style))
        if timestamp:
            story.append(Paragraph(timestamp, ts_style))
        story.append(Spacer(1, 0.2 * cm))

    doc_template.build(story)
    return buf.getvalue()
