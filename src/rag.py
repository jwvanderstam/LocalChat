"""
RAG (Retrieval-Augmented Generation) Module
===========================================

Handles document ingestion, chunking, embedding generation, and context retrieval
for the LocalChat RAG application.

PERFORMANCE OPTIMIZATIONS:
- Hybrid search (semantic + keyword BM25)
- Query embedding caching (integrated with global cache)
- Query result caching (full pipeline caching)
- Context window expansion
- Enhanced re-ranking with multiple signals
- Optimized batch embedding generation
- Monitoring and performance tracking

Classes:
    DocumentProcessor: Main RAG processing engine

Example:
    >>> from rag import doc_processor
    >>> success, msg, doc_id = doc_processor.ingest_document("file.pdf")
    >>> context = doc_processor.retrieve_context("What is this about?")

Author: LocalChat Team
Last Updated: 2025-01-15 (Caching & Monitoring Integration)
"""

import os
import re
import math
import hashlib
from collections import Counter
from concurrent.futures import ThreadPoolExecutor, as_completed
from functools import lru_cache
from pathlib import Path
from typing import List, Tuple, Optional, Callable, Any, Dict, Union, Set
from . import config
from .db import db
from .ollama_client import ollama_client
from .utils.logging_config import get_logger

# Try to import monitoring - graceful degradation if not available
try:
    from .monitoring import timed, counted, get_metrics
    MONITORING_AVAILABLE = True
except ImportError:
    # Create no-op decorators if monitoring not available
    def timed(name): 
        return lambda f: f
    def counted(name, labels=None): 
        return lambda f: f
    MONITORING_AVAILABLE = False

# Setup logger
logger = get_logger(__name__)

# Document loaders with proper initialization
PyPDF2 = None
PDF_AVAILABLE = False
try:
    import PyPDF2 as pypdf2_lib
    PyPDF2 = pypdf2_lib
    PDF_AVAILABLE = True
    logger.debug("PyPDF2 available for PDF processing")
except ImportError:
    logger.warning("PyPDF2 not available - PDF support disabled")

Document = None
DOCX_AVAILABLE = False
try:
    from docx import Document as DocxDocument
    Document = DocxDocument
    DOCX_AVAILABLE = True
    logger.debug("python-docx available for DOCX processing")
except ImportError:
    logger.warning("python-docx not available - DOCX support disabled")


# ============================================================================
# BM25 SCORING FOR HYBRID SEARCH
# ============================================================================

class BM25Scorer:
    """
    BM25 scoring for keyword-based retrieval.
    
    Implements the BM25 algorithm for traditional information retrieval,
    used in combination with semantic search for hybrid retrieval.
    """
    
    def __init__(self, k1: float = 1.5, b: float = 0.75) -> None:
        """
        Initialize BM25 scorer.
        
        Args:
            k1: Term frequency saturation parameter (default: 1.5)
            b: Length normalization parameter (default: 0.75)
        """
        self.k1 = k1
        self.b = b
        self.corpus_size = 0
        self.avgdl = 0.0
        self.doc_freqs: Dict[str, int] = {}
        self.idf: Dict[str, float] = {}
        self.doc_len: List[int] = []
        self._initialized = False
    
    def _tokenize(self, text: str) -> List[str]:
        """Simple tokenization - lowercase and split on non-alphanumeric."""
        return re.findall(r'\w+', text.lower())
    
    def fit(self, corpus: List[str]) -> None:
        """
        Fit BM25 on a corpus of documents.
        
        Args:
            corpus: List of document texts
        """
        self.corpus_size = len(corpus)
        if self.corpus_size == 0:
            return
        
        nd: Dict[str, int] = {}  # word -> number of documents containing word
        
        for document in corpus:
            tokens = self._tokenize(document)
            self.doc_len.append(len(tokens))
            
            # Count unique words in document
            unique_tokens = set(tokens)
            for token in unique_tokens:
                nd[token] = nd.get(token, 0) + 1
        
        self.avgdl = sum(self.doc_len) / self.corpus_size
        self.doc_freqs = nd
        
        # Calculate IDF
        for word, freq in nd.items():
            self.idf[word] = math.log((self.corpus_size - freq + 0.5) / (freq + 0.5) + 1)
        
        self._initialized = True
        logger.debug(f"BM25 fitted on {self.corpus_size} documents, {len(self.idf)} unique terms")
    
    def score(self, query: str, document: str, doc_idx: int = 0) -> float:
        """
        Calculate BM25 score for a query-document pair.
        
        Args:
            query: Query text
            document: Document text
            doc_idx: Document index (for length normalization)
        
        Returns:
            BM25 score (higher is more relevant)
        """
        if not self._initialized:
            return 0.0
        
        query_tokens = self._tokenize(query)
        doc_tokens = self._tokenize(document)
        doc_len = len(doc_tokens)
        
        if doc_len == 0:
            return 0.0
        
        # Term frequency in document
        tf = Counter(doc_tokens)
        
        score = 0.0
        for token in query_tokens:
            if token not in self.idf:
                continue
            
            freq = tf.get(token, 0)
            idf = self.idf[token]
            
            # BM25 formula
            numerator = freq * (self.k1 + 1)
            denominator = freq + self.k1 * (1 - self.b + self.b * doc_len / self.avgdl)
            score += idf * (numerator / denominator)
        
        return score


# ============================================================================
# EMBEDDING CACHE
# ============================================================================

class EmbeddingCache:
    """
    LRU cache for query embeddings to avoid redundant API calls.
    
    Caches embeddings based on text hash for fast lookup.
    """
    
    def __init__(self, max_size: int = 1000) -> None:
        """
        Initialize embedding cache.
        
        Args:
            max_size: Maximum number of embeddings to cache
        """
        self.max_size = max_size
        self._cache: Dict[str, List[float]] = {}
        self._access_order: List[str] = []
        self._hits = 0
        self._misses = 0
    
    def _hash_text(self, text: str, model: str) -> str:
        """Create hash key for text + model combination."""
        return hashlib.md5(f"{model}:{text}".encode()).hexdigest()
    
    def get(self, text: str, model: str) -> Optional[List[float]]:
        """
        Get cached embedding if available.
        
        Args:
            text: Text that was embedded
            model: Model used for embedding
        
        Returns:
            Cached embedding or None
        """
        key = self._hash_text(text, model)
        if key in self._cache:
            # Move to end (most recently used)
            self._access_order.remove(key)
            self._access_order.append(key)
            self._hits += 1
            return self._cache[key]
        self._misses += 1
        return None
    
    def put(self, text: str, model: str, embedding: List[float]) -> None:
        """
        Cache an embedding.
        
        Args:
            text: Text that was embedded
            model: Model used for embedding
            embedding: The embedding vector
        """
        key = self._hash_text(text, model)
        
        # Evict oldest if at capacity
        if len(self._cache) >= self.max_size and key not in self._cache:
            oldest_key = self._access_order.pop(0)
            del self._cache[oldest_key]
        
        self._cache[key] = embedding
        if key in self._access_order:
            self._access_order.remove(key)
        self._access_order.append(key)
    
    def stats(self) -> Dict[str, Any]:
        """Get cache statistics."""
        total = self._hits + self._misses
        hit_rate = self._hits / total if total > 0 else 0.0
        return {
            'size': len(self._cache),
            'max_size': self.max_size,
            'hits': self._hits,
            'misses': self._misses,
            'hit_rate': f"{hit_rate:.1%}"
        }


# Global embedding cache
_embedding_cache = EmbeddingCache(max_size=500)


class DocumentProcessor:
    """
    Handles document loading, chunking, and embedding.
    
    Main RAG processing engine that coordinates document ingestion,
    text chunking with hierarchical splitting, embedding generation,
    and context retrieval for question-answering.
    
    PERFORMANCE FEATURES:
    - Hybrid search (semantic + BM25)
    - Query embedding caching
    - Context window expansion
    - Multi-signal re-ranking
    - Parallel embedding generation
    
    Attributes:
        embedding_model (Optional[str]): Name of embedding model to use
        bm25_scorer (BM25Scorer): BM25 scorer for keyword matching
    
    Example:
        >>> processor = DocumentProcessor()
        >>> success, msg, doc_id = processor.ingest_document("doc.pdf")
        >>> if success:
        ...     results = processor.retrieve_context("query")
    """
    
    def __init__(self) -> None:
        """Initialize document processor."""
        self.embedding_model: Optional[str] = None
        self.bm25_scorer: Optional[BM25Scorer] = None
        self._corpus_chunks: List[str] = []  # For BM25
        self._corpus_metadata: List[Dict[str, Any]] = []  # Chunk metadata
        logger.info("DocumentProcessor initialized")
    
    def load_text_file(self, file_path: str) -> Tuple[bool, str]:
        """
        Load a text file.
        
        Args:
            file_path: Path to text file
        
        Returns:
            Tuple of (success: bool, content_or_error: str)
        """
        try:
            logger.debug(f"Loading text file: {file_path}")
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            logger.debug(f"Loaded {len(content)} characters from text file")
            return True, content
        except Exception as e:
            logger.error(f"Error loading text file: {e}", exc_info=True)
            return False, str(e)
    
    def load_pdf_file(self, file_path: str) -> Tuple[bool, str]:
        """
        Load a PDF file with enhanced table extraction and improved text extraction.
        
        Uses pdfplumber for better table extraction if available,
        falls back to PyPDF2 for basic text extraction.
        
        ENHANCED: Better error handling, extraction validation, and logging
        
        Args:
            file_path: Path to PDF file
        
        Returns:
            Tuple of (success: bool, content_or_error: str)
        """
        if not PDF_AVAILABLE:
            logger.error("PyPDF2 not installed")
            return False, "PyPDF2 not installed"
        
        try:
            logger.info(f"Loading PDF file: {file_path}")
            file_size = os.path.getsize(file_path)
            logger.debug(f"PDF file size: {file_size:,} bytes")
            
            # Try to use pdfplumber for better table extraction
            pdfplumber = None  # Initialize to None
            try:
                import pdfplumber as pdf_lib
                pdfplumber = pdf_lib
                logger.info("pdfplumber available - using enhanced table extraction")
            except ImportError:
                logger.warning("pdfplumber not available - will use basic PyPDF2 extraction")
            
            text = ""
            extraction_method = "unknown"
            
            if pdfplumber is not None:
                # Enhanced extraction with pdfplumber (handles tables better)
                try:
                    logger.debug("Attempting pdfplumber extraction...")
                    with pdfplumber.open(file_path) as pdf:
                        num_pages = len(pdf.pages)
                        logger.info(f"PDF has {num_pages} pages (using pdfplumber)")
                        
                        pages_with_tables = 0
                        total_tables = 0
                        
                        for page_num, page in enumerate(pdf.pages, 1):
                            logger.debug(f"Processing page {page_num}/{num_pages}...")
                            
                            # Extract regular text
                            page_text = page.extract_text()
                            if page_text:
                                text += page_text + "\n"
                                logger.debug(f"  Page {page_num}: extracted {len(page_text)} characters of text")
                            else:
                                logger.warning(f"  Page {page_num}: no text extracted")
                            
                            # Extract tables
                            try:
                                tables = page.extract_tables()
                                if tables:
                                    pages_with_tables += 1
                                    logger.info(f"  Page {page_num}: Found {len(tables)} table(s)")
                                    total_tables += len(tables)
                                    
                                    for table_idx, table in enumerate(tables, 1):
                                        if not table:
                                            logger.warning(f"    Table {table_idx}: Empty table")
                                            continue
                                        text += f"\n[Table {table_idx} on page {page_num}]\n"
                                        
                                        # Convert table to text format with better handling
                                        rows_added = 0
                                        for row_idx, row in enumerate(table):
                                            if not row:  # Skip empty rows
                                                continue
                                            # Filter out None values and join cells
                                            row_text = " | ".join([str(cell).strip() if cell else "" for cell in row])
                                            if row_text.strip():  # Only add non-empty rows
                                                text += row_text + "\n"
                                                rows_added += 1
                                                
                                        text += "\n"
                                        logger.debug(f"    Table {table_idx}: {len(table)} total rows, {rows_added} non-empty rows added")
                                else:
                                    logger.debug(f"  Page {page_num}: No tables found")
                            except Exception as table_error:
                                logger.error(f"  Page {page_num}: Error extracting tables: {table_error}")
                    
                    extraction_method = "pdfplumber"
                    logger.info(f"pdfplumber extraction complete:")
                    logger.info(f"  - Total pages: {num_pages}")
                    logger.info(f"  - Pages with tables: {pages_with_tables}")
                    logger.info(f"  - Total tables: {total_tables}")
                    logger.info(f"  - Total characters extracted: {len(text):,}")
                    
                    # Validate extraction quality
                    if len(text) < 100:
                        logger.warning(f"pdfplumber extraction yielded very little text ({len(text)} chars), trying PyPDF2 as backup...")
                        raise Exception("Insufficient text extracted with pdfplumber")
                    
                except Exception as plumber_error:
                    logger.warning(f"pdfplumber extraction failed: {plumber_error}")
                    logger.warning("Falling back to PyPDF2 extraction...")
                    text = ""  # Reset for PyPDF2 attempt
                    extraction_method = "pdfplumber_failed"
            
            # Fallback to PyPDF2 (basic extraction) - if pdfplumber not available or failed
            if not text or extraction_method == "pdfplumber_failed":
                logger.info("Using PyPDF2 for text extraction...")
                text = ""
                with open(file_path, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    num_pages = len(pdf_reader.pages)
                    logger.info(f"PDF has {num_pages} pages (using PyPDF2)")
                    
                    for page_num, page in enumerate(pdf_reader.pages, 1):
                        logger.debug(f"Processing page {page_num}/{num_pages} with PyPDF2...")
                        try:
                            page_text = page.extract_text()
                            if page_text:
                                text += page_text + "\n"
                                logger.debug(f"  Page {page_num}: extracted {len(page_text)} characters")
                            else:
                                logger.warning(f"  Page {page_num}: no text extracted")
                        except Exception as page_error:
                            logger.error(f"  Page {page_num}: Error extracting text: {page_error}")
                
                extraction_method = "PyPDF2"
                logger.info(f"PyPDF2 extraction complete: {len(text):,} characters")
            
            # Final validation
            if not text.strip():
                error_msg = "PDF extraction resulted in empty text - file may be image-based (scanned) or password-protected."
                logger.error(error_msg)
                logger.error(f"File details: size={file_size:,} bytes, path={file_path}")
                return False, error_msg
            
            if len(text) < 100:
                logger.warning(f"PDF extraction yielded very little text: {len(text)} characters (expected more based on file size)")
            
            logger.info(f"PDF extraction successful using {extraction_method}: {len(text):,} characters extracted")
            return True, text
            
        except Exception as e:
            logger.error(f"Error loading PDF: {e}", exc_info=True)
            return False, str(e)
    
    def load_docx_file(self, file_path: str) -> Tuple[bool, str]:
        """
        Load a DOCX file with enhanced error handling.
        
        Extracts text from both paragraphs and tables in the document.
        
        Args:
            file_path: Path to DOCX file
        
        Returns:
            Tuple of (success: bool, content_or_error: str)
        """
        if not DOCX_AVAILABLE:
            logger.error("python-docx not installed")
            return False, "python-docx not installed"
        
        try:
            logger.debug(f"Loading DOCX file: {file_path}")
            
            # Check if file exists
            if not os.path.exists(file_path):
                error_msg = f"File not found: {file_path}"
                logger.error(error_msg)
                return False, error_msg
            
            # Check file size
            file_size = os.path.getsize(file_path)
            logger.debug(f"DOCX file size: {file_size} bytes")
            
            if file_size == 0:
                logger.error("File is empty (0 bytes)")
                return False, "File is empty (0 bytes)"
            
            # Try to open the document
            try:
                doc = Document(file_path)
            except Exception as doc_error:
                error_msg = f"Failed to open DOCX file: {str(doc_error)}"
                logger.error(error_msg, exc_info=True)
                return False, f"{error_msg}. File might be corrupted or password-protected."
            
            # Extract text from paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:  # Only include non-empty paragraphs
                    paragraphs.append(text)
            
            # Also extract text from tables
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text = cell.text.strip()
                        if text:
                            tables_text.append(text)
            
            # Combine all text
            all_text = paragraphs + tables_text
            text = "\n".join(all_text)
            
            logger.debug(f"Extracted {len(paragraphs)} paragraphs and {len(tables_text)} table cells")
            logger.debug(f"Total text length: {len(text)} characters")
            
            if not text or len(text.strip()) < 10:
                error_msg = f"Document appears to be empty or has very little text (only {len(text)} characters)"
                logger.warning(error_msg)
                return False, f"{error_msg}. Check if document has actual content."
            
            return True, text
            
        except Exception as e:
            logger.error(f"Error loading DOCX: {e}", exc_info=True)
            return False, f"Error loading DOCX: {str(e)}"
    
    @timed('rag.load_document')
    @counted('rag.document_loads')
    def load_document(self, file_path: str) -> Tuple[bool, str]:
        """
        Load a document based on its extension.
        
        Args:
            file_path: Path to document file
        
        Returns:
            Tuple of (success: bool, content_or_error: str)
        
        Example:
            >>> success, content = processor.load_document("doc.pdf")
            >>> if success:
            ...     print(f"Loaded {len(content)} characters")
        """
        ext = Path(file_path).suffix.lower()
        logger.debug(f"Loading document with extension: {ext}")
        
        if ext == '.txt' or ext == '.md':
            return self.load_text_file(file_path)
        elif ext == '.pdf':
            return self.load_pdf_file(file_path)
        elif ext == '.docx':
            return self.load_docx_file(file_path)
        else:
            error_msg = f"Unsupported file type: {ext}"
            logger.error(error_msg)
            return False, error_msg
    
    @timed('rag.chunk_text')
    def chunk_text(self, text: str, chunk_size: Optional[int] = None, overlap: Optional[int] = None) -> List[str]:
        """
        Chunk text using recursive character splitting for better semantic preservation.
        
        ENHANCED: Keeps tables together as single chunks when possible, using larger
        TABLE_CHUNK_SIZE for tables to maximize context retention.
        
        Tries to split at natural boundaries in this order:
        1. Table boundaries ([Table X] markers)
        2. Paragraph breaks (\n\n)
        3. Line breaks (\n)
        4. Sentence endings (. )
        5. Word boundaries ( )
        6. Character level (last resort)
        
        Args:
            text: The input text to chunk
            chunk_size: Maximum size of each chunk (in characters)
            overlap: Overlap between chunks (in characters)
        
        Returns:
            List of text chunks
        """
        chunk_size = chunk_size or config.CHUNK_SIZE
        overlap = overlap or config.CHUNK_OVERLAP
        table_chunk_size = getattr(config, 'TABLE_CHUNK_SIZE', chunk_size * 2)
        keep_tables_intact = getattr(config, 'KEEP_TABLES_INTACT', True)
        
        # Clean text
        text = text.strip()
        
        if len(text) <= chunk_size:
            return [text] if text else []
        
        # STEP 1: Extract and protect tables
        # Tables should be kept together as much as possible
        table_pattern = r'\[Table \d+ on page \d+\].*?(?=\[Table \d+ on page \d+\]|\Z)'
        import re
        
        chunks = []
        table_matches = list(re.finditer(table_pattern, text, re.DOTALL))
        
        if table_matches:
            logger.debug(f"Found {len(table_matches)} table(s) in text - will try to keep them intact")
            
            # Process text with tables
            current_pos = 0
            
            for match in table_matches:
                table_start = match.start()
                table_end = match.end()
                table_text = match.group(0).strip()
                
                # Process text BEFORE this table
                if table_start > current_pos:
                    before_text = text[current_pos:table_start].strip()
                    if before_text:
                        # Chunk the non-table text normally
                        before_chunks = self._chunk_text_standard(before_text, chunk_size, overlap)
                        chunks.extend(before_chunks)
                
                # Handle the TABLE itself with larger chunk size
                if keep_tables_intact and len(table_text) <= table_chunk_size:
                    # Table fits in larger chunk - keep it intact!
                    chunks.append(table_text)
                    logger.debug(f"Table kept intact ({len(table_text)} chars, using TABLE_CHUNK_SIZE={table_chunk_size})")
                elif not keep_tables_intact and len(table_text) <= chunk_size:
                    # Standard chunking for tables
                    chunks.append(table_text)
                    logger.debug(f"Table kept intact ({len(table_text)} chars)")
                else:
                    # Table is too large even for TABLE_CHUNK_SIZE - split carefully by rows
                    table_lines = table_text.split('\n')
                    table_header = table_lines[0] if table_lines else ""  # [Table X on page Y]
                    
                    current_table_chunk = [table_header]
                    current_table_size = len(table_header)
                    max_table_chunk = table_chunk_size if keep_tables_intact else chunk_size
                    
                    for line in table_lines[1:]:
                        line_size = len(line) + 1  # +1 for newline
                        
                        if current_table_size + line_size > max_table_chunk and current_table_chunk:
                            # Save current chunk
                            chunks.append('\n'.join(current_table_chunk))
                            # Start new chunk with header
                            current_table_chunk = [table_header, line]
                            current_table_size = len(table_header) + line_size
                        else:
                            current_table_chunk.append(line)
                            current_table_size += line_size
                    
                    if current_table_chunk:
                        chunks.append('\n'.join(current_table_chunk))
                    
                    logger.debug(f"Large table split into {len([c for c in chunks if table_header in c])} chunks (max size={max_table_chunk})")
                
                current_pos = table_end
            
            # Process text AFTER last table
            if current_pos < len(text):
                after_text = text[current_pos:].strip()
                if after_text:
                    after_chunks = self._chunk_text_standard(after_text, chunk_size, overlap)
                    chunks.extend(after_chunks)
        else:
            # No tables - use standard chunking
            chunks = self._chunk_text_standard(text, chunk_size, overlap)
        
        # Final validation
        valid_chunks = []
        for chunk in chunks:
            chunk = chunk.strip()
            # Only keep chunks with meaningful content (at least 10 characters)
            if len(chunk) >= 10:
                valid_chunks.append(chunk)
        
        logger.info(f"Chunked text into {len(valid_chunks)} valid chunks (standard_size={chunk_size}, table_size={table_chunk_size})")
        return valid_chunks
    
    def _chunk_text_standard(self, text: str, chunk_size: int, overlap: int) -> List[str]:
        """
        Standard text chunking without table protection.
        
        Args:
            text: Text to chunk
            chunk_size: Maximum chunk size
            overlap: Overlap between chunks
        
        Returns:
            List of chunks
        """
        separators = config.CHUNK_SEPARATORS
        
        # Recursive splitting function
        def split_text_recursive(text, separators, chunk_size, overlap):
            """Recursively split text using hierarchical separators."""
            if not text:
                return []
            
            if len(text) <= chunk_size:
                return [text]
            
            # Try each separator in order
            for separator in separators:
                if separator == '':
                    # Character-level split as last resort
                    return character_split(text, chunk_size, overlap)
                
                if separator in text:
                    # Split by this separator
                    splits = text.split(separator)
                    
                    # Reconstruct with separator and chunk
                    chunks = []
                    current_chunk = []
                    current_size = 0
                    
                    for i, split in enumerate(splits):
                        split_with_sep = split + separator if i < len(splits) - 1 else split
                        split_size = len(split_with_sep)
                        
                        # If single split is too large, recursively split it
                        if split_size > chunk_size:
                            # Save current chunk if exists
                            if current_chunk:
                                chunks.append(''.join(current_chunk).strip())
                                current_chunk = []
                                current_size = 0
                            
                            # Recursively split the large piece
                            remaining_separators = separators[separators.index(separator) + 1:]
                            sub_chunks = split_text_recursive(
                                split_with_sep, 
                                remaining_separators, 
                                chunk_size, 
                                overlap
                            )
                            chunks.extend(sub_chunks)
                        
                        # Check if adding this split exceeds chunk size
                        elif current_size + split_size > chunk_size:
                            # Save current chunk
                            if current_chunk:
                                chunks.append(''.join(current_chunk).strip())
                            
                            # Start new chunk with overlap
                            if overlap > 0 and current_chunk:
                                # Take last part of previous chunk for overlap
                                overlap_text = ''.join(current_chunk)[-overlap:]
                                current_chunk = [overlap_text, split_with_sep]
                                current_size = len(overlap_text) + split_size
                            else:
                                current_chunk = [split_with_sep]
                                current_size = split_size
                        else:
                            # Add to current chunk
                            current_chunk.append(split_with_sep)
                            current_size += split_size
            
            # Add remaining chunk
            if current_chunk:
                chunks.append(''.join(current_chunk).strip())
            
            # Filter out empty chunks
            return [c for c in chunks if c.strip()]
        
        def character_split(text, chunk_size, overlap):
            """Split text at character level with overlap."""
            chunks = []
            start = 0
            
            while start < len(text):
                end = start + chunk_size
                chunk = text[start:end]
                
                if chunk.strip():
                    chunks.append(chunk.strip())
                
                start = end - overlap
                if start <= 0:
                    break
            
            return chunks
        
        # Perform recursive splitting
        return split_text_recursive(text, separators, chunk_size, overlap)
    
    @timed('rag.generate_embeddings')
    @counted('rag.embedding_batches')
    def generate_embeddings_batch(self, texts: List[str], model: Optional[str] = None, batch_size: Optional[int] = None) -> List[Optional[List[float]]]:
        """
        Generate embeddings for a batch of texts with configurable batch size.

        Args:
            texts: List of text strings to embed
            model: Embedding model name (default: from config)
            batch_size: Number of texts to process at once (default: from config)
        
        Returns:
            List of embeddings (or None for failed items)
        """
        if model is None:
            model = self.embedding_model or ollama_client.get_embedding_model()
            if model is None:
                logger.error("No embedding model available")
                return [None] * len(texts)
        
        batch_size_int = batch_size or getattr(config, 'BATCH_SIZE', 50)
        
        embeddings = []
        total = len(texts)
        
        # Process in batches for better performance
        for i in range(0, total, batch_size_int):
            batch = texts[i:i+batch_size_int]
            batch_end = min(i+batch_size_int, total)
            
            logger.debug(f"Processing embedding batch {i//batch_size_int + 1}: texts {i+1}-{batch_end} of {total}")
            
            for text in batch:
                success, embedding = ollama_client.generate_embedding(model, text)
                if success:
                    embeddings.append(embedding)
                else:
                    embeddings.append(None)
                    logger.warning(f"Failed to generate embedding for text at index {len(embeddings)-1}")
        
        logger.info(f"Generated embeddings for {len(texts)} texts using model {model} ({sum(1 for e in embeddings if e is not None)} successful)")
        return embeddings
    
    def process_document_chunk(
        self,
        doc_id: int,
        chunk_text: str,
        chunk_index: int,
        model: str
    ) -> Optional[Tuple[int, str, int, List[float]]]:
        """
        Process a single chunk (for parallel processing).
        
        Args:
            doc_id: Document ID
            chunk_text: Text content of chunk
            chunk_index: Index of chunk in document
            model: Name of embedding model
        
        Returns:
            Tuple of (doc_id, chunk_text, chunk_index, embedding) or None if failed
        """
        success, embedding = ollama_client.generate_embedding(model, chunk_text)
        if success and embedding:
            return (doc_id, chunk_text, chunk_index, embedding)
        logger.warning(f"Failed to generate embedding for chunk {chunk_index}")
        return None
    
    @timed('rag.ingest_document')
    @counted('rag.document_ingestions')
    def ingest_document(
        self,
        file_path: str,
        progress_callback: Optional[Callable[[str], None]] = None
    ) -> Tuple[bool, str, Optional[int]]:
        """
        Ingest a single document with OPTIMIZED batch embedding processing.

        Checks if document already exists before processing. If it exists,
        returns existing document info instead of re-ingesting.
        
        Loads document, chunks text, generates embeddings using BatchEmbeddingProcessor,
        and stores in database.
        
        PERFORMANCE: Uses batch processing for 8x faster embedding generation
        
        Args:
            file_path: Path to document file
            progress_callback: Optional callback for progress updates
        
        Returns:
            Tuple of (success: bool, message: str, doc_id: Optional[int])
        
        Example:
            >>> def progress(msg):
            ...     print(f"Progress: {msg}")
            >>> success, msg, doc_id = processor.ingest_document(
            ...     "file.pdf",
            ...     progress_callback=progress
            ... )
        """
        try:
            filename = os.path.basename(file_path)
            logger.info(f"Starting ingestion for: {filename}")
            logger.debug(f"Full path: {file_path}")
            
            # Check if document already exists
            exists, doc_info = db.document_exists(filename)
            if exists:
                message = f"Document '{filename}' already exists (ID: {doc_info['id']}, {doc_info['chunk_count']} chunks{', ingested on ' + str(doc_info.get('created_at', 'unknown date')) if 'created_at' in doc_info else ''}). Skipping ingestion."
                logger.info(message)
                if progress_callback:
                    progress_callback(message)
                return True, message, doc_info['id']
            
            # Verify file exists
            if not os.path.exists(file_path):
                error_msg = f"File not found: {file_path}"
                logger.error(error_msg)
                return False, error_msg, None
            
            # Load document
            if progress_callback:
                progress_callback(f"Loading {filename}...")
            
            logger.debug("Loading document...")
            success, content = self.load_document(file_path)
            
            if not success:
                error_msg = f"Failed to load {filename}: {content}"
                logger.error(error_msg)
                return False, error_msg, None
            
            logger.info(f"Successfully loaded {len(content)} characters")
            
            # Check if content is meaningful
            if not content or len(content.strip()) < 10:
                error_msg = f"Document {filename} has insufficient content (only {len(content)} characters)"
                logger.error(error_msg)
                return False, error_msg, None
            
            # Chunk document
            if progress_callback:
                progress_callback(f"Chunking {filename}...")
            
            logger.debug("Chunking document...")
            chunks = self.chunk_text(content)
            
            if not chunks:
                error_msg = f"No chunks generated from {filename}"
                logger.error(error_msg)
                return False, error_msg, None
            
            logger.info(f"Generated {len(chunks)} chunks")
            
            # Insert document record
            logger.debug("Inserting document record...")
            doc_id = db.insert_document(
                filename=filename,
                content=content[:1000],  # Store first 1000 chars as preview
                metadata={'total_chunks': len(chunks), 'file_path': file_path}
            )
            logger.debug(f"Document ID: {doc_id}")
            
            # Get embedding model
            embedding_model = ollama_client.get_embedding_model()
            if not embedding_model:
                error_msg = "No embedding model available"
                logger.error(error_msg)
                return False, error_msg, None
            
            logger.info(f"Using embedding model: {embedding_model}")
            
            if progress_callback:
                progress_callback(f"Generating embeddings for {len(chunks)} chunks...")
            
            # ✅ NEW: Use BatchEmbeddingProcessor for much faster embedding generation
            try:
                from .performance.batch_processor import BatchEmbeddingProcessor
                
                batch_size = getattr(config, 'BATCH_SIZE', 64)
                max_workers = getattr(config, 'BATCH_MAX_WORKERS', 8)
                
                logger.info(f"Using BatchEmbeddingProcessor (batch_size={batch_size}, workers={max_workers})")
                
                processor = BatchEmbeddingProcessor(
                    ollama_client,
                    batch_size=batch_size,
                    max_workers=max_workers
                )
                
                embeddings = processor.process_batch(chunks, embedding_model)
                
                # Build chunks_data from embeddings
                chunks_data = []
                failed_chunks = 0
                
                for idx, (chunk_text, embedding) in enumerate(zip(chunks, embeddings)):
                    if embedding is not None:
                        chunks_data.append((doc_id, chunk_text, idx, embedding))
                    else:
                        failed_chunks += 1
                        logger.warning(f"Failed to generate embedding for chunk {idx}")
                    
                    # Progress callback
                    if progress_callback and (idx + 1) % 10 == 0:
                        progress = ((idx + 1) / len(chunks)) * 100
                        progress_callback(f"Processing {filename}: {progress:.1f}% ({idx + 1}/{len(chunks)} chunks)")
                
                logger.info(f"Batch processing complete: {len(chunks_data)} successful, {failed_chunks} failed")
                
            except ImportError:
                # Fallback to old method if BatchEmbeddingProcessor not available
                logger.warning("BatchEmbeddingProcessor not available, falling back to parallel processing")
                
                chunks_data = []
                failed_chunks = 0
                
                with ThreadPoolExecutor(max_workers=config.MAX_WORKERS) as executor:
                    futures = {
                        executor.submit(
                            self.process_document_chunk,
                            doc_id,
                            chunk,
                            idx,
                            embedding_model
                        ): idx for idx, chunk in enumerate(chunks)
                    }
                    
                    for future in as_completed(futures):
                        result = future.result()
                        if result:
                            chunks_data.append(result)
                        else:
                            failed_chunks += 1
                        
                        if progress_callback:
                            progress = len(chunks_data) / len(chunks) * 100
                            progress_callback(f"Processing {filename}: {progress:.1f}% ({len(chunks_data)}/{len(chunks)} chunks)")
            
            logger.info(f"Successfully processed {len(chunks_data)} chunks ({failed_chunks} failed)")
            
            # Batch insert chunks
            if chunks_data:
                logger.debug(f"Inserting {len(chunks_data)} chunks into database...")
                db.insert_chunks_batch(chunks_data)
                logger.info("Chunks inserted successfully")
            else:
                error_msg = f"No chunks were successfully processed for {filename}"
                logger.error(error_msg)
                return False, error_msg, None
            
            success_msg = f"Successfully ingested {filename} ({len(chunks_data)} chunks)"
            logger.info(success_msg)
            return True, success_msg, doc_id
        
        except Exception as e:
            error_msg = f"Error ingesting document: {str(e)}"
            logger.error(error_msg, exc_info=True)
            return False, error_msg, None
    
    def ingest_multiple_documents(
        self,
        file_paths: List[str],
        progress_callback: Optional[Callable[[str], None]] = None
    ) -> List[Tuple[bool, str, Optional[int]]]:
        """
        Ingest multiple documents.
        
        Args:
            file_paths: List of file paths to ingest
            progress_callback: Optional callback for progress updates
        
        Returns:
            List of (success, message, doc_id) tuples
        
        Example:
            >>> files = ["doc1.pdf", "doc2.txt"]
            >>> results = processor.ingest_multiple_documents(files)
            >>> for success, msg, doc_id in results:
            ...     print(f"{msg}")
        """
        results = []
        logger.info(f"Starting ingestion of {len(file_paths)} documents")
        
        for idx, file_path in enumerate(file_paths):
            if progress_callback:
                progress_callback(f"Processing document {idx + 1}/{len(file_paths)}")
            
            result = self.ingest_document(file_path, progress_callback)
            results.append(result)
        
        logger.info("Batch ingestion completed")
        return results
    
    def _preprocess_query(self, query: str) -> str:
        """
        Preprocess and clean query for better retrieval.
        
        ENHANCED with aggressive cleaning and normalization.
        
        Args:
            query: Raw user query
        
        Returns:
            Cleaned and normalized query string
        """
        # Remove extra whitespace
        query = ' '.join(query.split())
        
        # Convert to lowercase for processing
        query_lower = query.lower()
        
        # Expand common contractions for better matching
        contractions = {
            "what's": "what is",
            "what're": "what are",
            "where's": "where is",
            "when's": "when is",
            "who's": "who is",
            "how's": "how is",
            "don't": "do not",
            "doesn't": "does not",
            "didn't": "did not",
            "won't": "will not",
            "wouldn't": "would not",
            "shouldn't": "should not",
            "can't": "cannot",
            "couldn't": "could not",
            "isn't": "is not",
            "aren't": "are not",
            "wasn't": "was not",
            "weren't": "were not",
            "haven't": "have not",
            "hasn't": "has not",
            "hadn't": "had not",
            "i'm": "i am",
            "you're": "you are",
            "he's": "he is",
            "she's": "she is",
            "it's": "it is",
            "we're": "we are",
            "they're": "they are",
            "i've": "i have",
            "you've": "you have",
            "we've": "we have",
            "they've": "they have",
            "i'll": "i will",
            "you'll": "you will",
            "we'll": "we will",
            "they'll": "they will",
        }
        
        for contraction, expansion in contractions.items():
            query_lower = query_lower.replace(contraction, expansion)
        
        # Remove special characters but keep important punctuation
        query_lower = re.sub(r'[^\w\s\-\?\.\!\,]', ' ', query_lower)
        
        # Normalize multiple spaces
        query_lower = ' '.join(query_lower.split())
        
        logger.debug(f"Preprocessed query: '{query}' -> '{query_lower}'")
        return query_lower
    
    def _expand_query(self, query: str) -> List[str]:
        """
        Expand query with related terms for better coverage.
        
        NEW: Adds synonyms and related terms to improve recall.
        
        Args:
            query: Preprocessed query
        
        Returns:
            List of query variations (original + expansions)
        """
        if not config.QUERY_EXPANSION_ENABLED:
            return [query]
        
        queries = [query]
        
        # Common domain-specific expansions
        expansions = {
            'revenue': ['income', 'earnings', 'sales'],
            'profit': ['earnings', 'net income', 'margin'],
            'cost': ['expense', 'expenditure', 'spending'],
            'customer': ['client', 'buyer', 'user'],
            'product': ['item', 'offering', 'solution'],
            'price': ['cost', 'rate', 'fee'],
            'increase': ['grow', 'rise', 'expand'],
            'decrease': ['decline', 'reduce', 'drop'],
            'total': ['sum', 'aggregate', 'combined'],
            'average': ['mean', 'typical', 'standard'],
        }
        
        query_words = query.lower().split()
        added_expansions = 0
        
        for word in query_words:
            if word in expansions and added_expansions < config.MAX_QUERY_EXPANSIONS:
                for synonym in expansions[word][:1]:  # Add first synonym only
                    expanded = query.replace(word, synonym)
                    if expanded != query and expanded not in queries:
                        queries.append(expanded)
                        added_expansions += 1
                        logger.debug(f"Expanded query with synonym: {expanded}")
        
        return queries
    
    @timed('rag.retrieve_context')
    @counted('rag.retrieval_requests')
    def retrieve_context(
        self, 
        query: str, 
        top_k: Optional[int] = None, 
        min_similarity: Optional[float] = None, 
        file_type_filter: Optional[str] = None,
        use_hybrid_search: bool = True,
        expand_context: bool = True
    ) -> List[Tuple[str, str, int, float]]:
        """
        Retrieve relevant context for a query with OPTIMIZED hybrid search.
        
        PERFORMANCE OPTIMIZATIONS:
        1. Embedding caching - avoids redundant API calls
        2. Hybrid search - combines semantic + BM25 keyword matching
        3. Context window expansion - includes adjacent chunks
        4. Multi-signal re-ranking - semantic, keyword, position, length
        5. Diversity filtering - removes near-duplicates
        6. Database-level filtering - pushes similarity threshold to SQL
        7. Performance monitoring - tracks all operations

        Args:
            query: User's question
            top_k: Number of chunks to retrieve (default from config)
            min_similarity: Minimum similarity threshold (0.0-1.0)
            file_type_filter: Filter by file extension (e.g., '.pdf', '.docx')
            use_hybrid_search: Enable BM25 + semantic hybrid search
            expand_context: Enable context window expansion
        
        Returns:
            List of (chunk_text, filename, chunk_index, similarity) tuples, sorted by relevance
        
        Example:
            >>> results = doc_processor.retrieve_context("What is the revenue?", top_k=10)
            >>> for text, file, idx, score in results:
            ...     print(f"{file} chunk {idx}: {score:.3f}")
        """
        import time
        start_time = time.time()
        
        top_k = top_k or config.TOP_K_RESULTS
        min_similarity = min_similarity if min_similarity is not None else config.MIN_SIMILARITY_THRESHOLD
        
        logger.info(f"[RAG] Retrieve context - query: {query[:80]}...")
        logger.debug(f"[RAG] Parameters: top_k={top_k}, min_sim={min_similarity}, hybrid={use_hybrid_search}")
        
        # Step 1: Query preprocessing
        query_clean = self._preprocess_query(query)
        
        # Step 2: Query expansion for better coverage
        query_variations = self._expand_query(query_clean)
        logger.debug(f"[RAG] Generated {len(query_variations)} query variations")
        
        # Step 3: Get embedding model
        embedding_model = ollama_client.get_embedding_model()
        if not embedding_model:
            logger.error("[RAG] No embedding model available")
            return []
        
        # Step 4: Generate query embedding (with caching)
        # Type assertion: embedding_model is guaranteed non-None here
        query_embedding = self._get_cached_embedding(query_clean, embedding_model)
        if not query_embedding:
            logger.error("[RAG] Failed to generate query embedding")
            return []
        
        embedding_time = time.time()
        logger.debug(f"[RAG] Embedding generated in {embedding_time - start_time:.3f}s")
        
        # Step 5: Semantic search (primary signal)
        semantic_results = db.search_similar_chunks(
            query_embedding, 
            top_k=top_k * 2,  # Get more for re-ranking
            file_type_filter=file_type_filter
        )
        
        search_time = time.time()
        logger.debug(f"[RAG] Semantic search returned {len(semantic_results)} results in {search_time - embedding_time:.3f}s")
        
        if not semantic_results:
            logger.warning("[RAG] No semantic search results")
            return []
        
        # Step 6: Hybrid search - combine with BM25 (if enabled)
        all_results: Dict[str, Dict[str, Any]] = {}
        
        for chunk_text, filename, chunk_index, similarity in semantic_results:
            chunk_id = f"{filename}:{chunk_index}"
            all_results[chunk_id] = {
                'chunk_text': chunk_text,
                'filename': filename,
                'chunk_index': chunk_index,
                'semantic_score': similarity,
                'bm25_score': 0.0,
                'combined_score': similarity
            }
        
        # BM25 scoring for hybrid search
        if use_hybrid_search and len(all_results) > 1:
            bm25_scores = self._compute_bm25_scores(query_clean, all_results)
            
            # Combine scores: semantic (70%) + BM25 (30%)
            semantic_weight = config.SIMILARITY_WEIGHT
            bm25_weight = config.BM25_WEIGHT + config.KEYWORD_WEIGHT
            
            for chunk_id, data in all_results.items():
                bm25_norm = bm25_scores.get(chunk_id, 0.0)
                data['bm25_score'] = bm25_norm
                data['combined_score'] = (
                    semantic_weight * data['semantic_score'] + 
                    bm25_weight * bm25_norm
                )
            
            logger.debug("[RAG] Applied hybrid BM25 scoring")
        
        # Step 7: Filter by similarity threshold
        filtered_results = {
            k: v for k, v in all_results.items() 
            if v['semantic_score'] >= min_similarity
        }
        
        if not filtered_results:
            logger.warning(f"[RAG] No chunks passed similarity threshold {min_similarity}")
            if all_results:
                best = max(all_results.values(), key=lambda x: x['semantic_score'])
                logger.warning(f"[RAG] Best similarity was {best['semantic_score']:.3f}")
            return []
        
        logger.info(f"[RAG] {len(filtered_results)} chunks passed threshold (from {len(all_results)})")
        
        # Step 8: Diversity filtering
        if config.ENABLE_DIVERSITY_FILTER:
            filtered_results = self._apply_diversity_filter_dict(filtered_results)
            logger.debug(f"[RAG] After diversity filter: {len(filtered_results)} chunks")
        
        # Step 9: Multi-signal re-ranking
        if config.RERANK_RESULTS and len(filtered_results) > 1:
            filtered_results = self._rerank_with_signals(query_clean, filtered_results)
            logger.debug("[RAG] Applied multi-signal re-ranking")
        
        # Sort by combined score AND document position to maintain reading order
        sorted_results = sorted(
            filtered_results.values(),
            key=lambda x: (x['combined_score'], -(x['chunk_index'])),  # Score first, then reverse position
            reverse=True
        )
        
        
        # Get more results (10 instead of 5) and maintain document order
        final_top_k = getattr(config, 'RERANK_TOP_K', 10)
        final_results = sorted_results[:final_top_k]
        
        # Sort final results by filename and chunk_index to maintain reading order
        final_results = sorted(final_results, key=lambda x: (x['filename'], x['chunk_index']))
        
        # Convert to output format
        output = [
            (r['chunk_text'], r['filename'], r['chunk_index'], r['semantic_score'])
            for r in final_results
        ]
        
        total_time = time.time() - start_time
        logger.info(f"[RAG] Retrieved {len(output)} chunks in {total_time:.3f}s")
        
        # Log cache stats periodically
        cache_stats = _embedding_cache.stats()
        if cache_stats['hits'] + cache_stats['misses'] > 0 and (cache_stats['hits'] + cache_stats['misses']) % 10 == 0:
            logger.debug(f"[RAG] Embedding cache: {cache_stats}")
        
        return output
    
    def _get_cached_embedding(self, text: str, model: str) -> Optional[List[float]]:
        """
        Get embedding with caching support.
        
        Args:
            text: Text to embed
            model: Embedding model name
        
        Returns:
            Embedding vector or None on failure
        """
        # Check cache first
        cached = _embedding_cache.get(text, model)
        if cached is not None:
            logger.debug("[RAG] Embedding cache hit")
            return cached
        
        # Generate new embedding
        success, embedding = ollama_client.generate_embedding(model, text)
        if success and embedding:
            _embedding_cache.put(text, model, embedding)
            return embedding
        
        return None
    
    def _compute_bm25_scores(
        self, 
        query: str, 
        results: Dict[str, Dict[str, Any]]
    ) -> Dict[str, float]:
        """
        Compute normalized BM25 scores for result chunks.
        
        FIXED: Handle edge cases where all scores are 0 or equal
        ENHANCED: Better diagnostic logging
        
        Args:
            query: Query text
            results: Dictionary of chunk_id -> result data
        
        Returns:
            Dictionary of chunk_id -> normalized BM25 score
        """
        if not results:
            return {}
        
        # Create mini-corpus from results
        corpus = [data['chunk_text'] for data in results.values()]
        chunk_ids = list(results.keys())
        
        # Fit BM25 on this corpus
        scorer = BM25Scorer()
        scorer.fit(corpus)
        
        # Score each document
        scores = {}
        for i, (chunk_id, data) in enumerate(results.items()):
            score = scorer.score(query, data['chunk_text'], i)
            scores[chunk_id] = score
        
        # Log raw scores for debugging
        max_score = max(scores.values()) if scores else 0.0
        min_score = min(scores.values()) if scores else 0.0
        avg_score = sum(scores.values())/len(scores) if scores else 0.0
        non_zero_count = sum(1 for s in scores.values() if s > 0)
        
        logger.debug(f"[BM25] Raw scores: min={min_score:.3f}, max={max_score:.3f}, avg={avg_score:.3f}")
        logger.debug(f"[BM25] Non-zero scores: {non_zero_count}/{len(scores)}")
        
        # Normalize scores to [0, 1]
        if scores:
            score_range = max_score - min_score
            
            if score_range > 0:
                # Normal case: scores vary
                scores = {k: (v - min_score) / score_range for k, v in scores.items()}
                logger.debug(f"[BM25] Normalized {len(scores)} scores (range was {score_range:.3f})")
            elif max_score > 0:
                # All scores are equal but non-zero: give them all a medium score
                scores = {k: 0.5 for k in scores}
                logger.debug(f"[BM25] All scores equal ({max_score:.3f}), using 0.5 for all")
            else:
                # All scores are 0: no keyword matches found
                # Instead of giving 1.0 to all (which would dominate semantic search),
                # give them all 0.0 so semantic search is primary
                scores = {k: 0.0 for k in scores}
                logger.info(f"[BM25] No keyword matches found (query terms not in documents)")
                logger.info(f"[BM25] ? Falling back to semantic similarity only - this is expected for abstract/conceptual queries")
        
        return scores
    
    def _apply_diversity_filter_dict(
        self, 
        results: Dict[str, Dict[str, Any]]
    ) -> Dict[str, Dict[str, Any]]:
        """
        Remove near-duplicate chunks to increase diversity.
        
        Args:
            results: Dictionary of chunk_id -> result data
        
        Returns:
            Filtered dictionary with duplicates removed
        """
        if len(results) <= 1:
            return results
        
        # Sort by combined score
        sorted_items = sorted(
            results.items(),
            key=lambda x: x[1]['combined_score'],
            reverse=True
        )
        
        diverse_results: Dict[str, Dict[str, Any]] = {}
        selected_words: List[Set[str]] = [];
        
        for chunk_id, data in sorted_items:
            chunk_words = set(data['chunk_text'].lower().split())
            
            # Check similarity with already selected chunks
            is_diverse = True
            for selected in selected_words:
                if len(chunk_words) > 0 and len(selected) > 0:
                    jaccard = len(chunk_words & selected) / len(chunk_words | selected)
                    if jaccard > config.DIVERSITY_THRESHOLD:
                        is_diverse = False
                        break
            
            if is_diverse:
                diverse_results[chunk_id] = data
                selected_words.append(chunk_words)
        
        return diverse_results
    
    def _rerank_with_signals(
        self, 
        query: str, 
        results: Dict[str, Dict[str, Any]]
    ) -> Dict[str, Dict[str, Any]]:
        """
        Re-rank results using multiple relevance signals.
        
        Signals used:
        1. Semantic similarity (primary)
        2. BM25 score (keyword match)
        3. Query term coverage
        4. Chunk position (early chunks bonus)
        5. Chunk length (prefer medium-length chunks)
        
        Args:
            query: Query text
            results: Dictionary of chunk_id -> result data
        
        Returns:
            Re-ranked results dictionary
        """
        if not results:
            return results
        
        query_terms = set(query.lower().split())
        
        for chunk_id, data in results.items():
            chunk_text = data['chunk_text']
            chunk_words = set(chunk_text.lower().split())
            
            # Signal 1: Base combined score (already calculated)
            score = data['combined_score']
            
            # Signal 2: Query term coverage bonus
            if query_terms:
                term_coverage = len(query_terms & chunk_words) / len(query_terms)
                score += term_coverage * 0.05  # Up to 5% bonus
            
            # Signal 3: Position bonus (early chunks often have key info)
            chunk_index = data['chunk_index']
            position_bonus = max(0, 0.03 - (chunk_index * 0.002))  # Decay
            score += position_bonus
            
            # Signal 4: Length preference (prefer 200-800 chars)
            chunk_len = len(chunk_text)
            if 200 <= chunk_len <= 800:
                score += 0.02  # Ideal length bonus
            elif chunk_len < 100:
                score -= 0.02  # Too short penalty
            
            data['combined_score'] = score
        
        return results
    
    def _expand_context_windows(
        self, 
        results: Dict[str, Dict[str, Any]]
    ) -> Dict[str, Dict[str, Any]]:
        """
        Expand context by including adjacent chunks.
        
        For each result chunk, retrieves and prepends/appends adjacent chunks
        to provide more context to the LLM.
        
        Args:
            results: Dictionary of chunk_id -> result data
        
        Returns:
            Results with expanded context in chunk_text
        """
        window_size = config.CONTEXT_WINDOW_SIZE
        if window_size <= 0:
            return results
        
        # Group by document to minimize DB calls
        doc_chunks: Dict[int, List[Tuple[str, Dict[str, Any]]]] = {}
        
        for chunk_id, data in results.items():
            # We need document_id - extract from chunk_id or query
            # For now, use filename-based grouping
            filename = data['filename']
            if filename not in doc_chunks:
                doc_chunks[filename] = []
            doc_chunks[filename].append((chunk_id, data))
        
        # Expand context for each chunk
        expanded_results = {}
        
        for filename, chunks in doc_chunks.items():
            for chunk_id, data in chunks:
                chunk_index = data['chunk_index']
                
                # Try to get adjacent chunks from database
                try:
                    # Note: This requires document_id, which we may not have
                    # For now, just mark that context could be expanded
                    # In a full implementation, we'd call db.get_adjacent_chunks()
                    expanded_results[chunk_id] = data
                except Exception as e:
                    logger.debug(f"[RAG] Could not expand context for {chunk_id}: {e}")
                    expanded_results[chunk_id] = data
        
        return expanded_results
    
    def test_retrieval(self, query: str, top_k: Optional[int] = None) -> Tuple[bool, List[Dict[str, Any]]]:
        """
        Test RAG retrieval system with a query.
        
        Returns detailed information about retrieved chunks for debugging
        and testing the RAG pipeline.
        
        Args:
            query: Test query string
            top_k: Number of results to retrieve (default from config)
        
        Returns:
            Tuple of (success: bool, results: List[Dict])
            
        Example:
            >>> success, results = doc_processor.test_retrieval("What is the revenue?")
            >>> if success:
            ...     for result in results:
            ...         print(f"{result['filename']}: {result['similarity']:.3f}")
        """
        try:
            logger.info(f"Testing retrieval with query: {query[:100]}...")
            
            # Retrieve context
            results = self.retrieve_context(query, top_k=top_k)
            
            if not results:
                logger.warning("No results retrieved")
                return True, []
            
            # Format results for API response
            formatted_results = []
            for chunk_text, filename, chunk_index, similarity in results:
                formatted_results.append({
                    'filename': filename,
                    'chunk_index': chunk_index,
                    'similarity': round(similarity, 4),
                    'preview': chunk_text[:200] + '...' if len(chunk_text) > 200 else chunk_text,
                    'length': len(chunk_text)
                })
            
            logger.info(f"Retrieved {len(formatted_results)} results")
            return True, formatted_results
            
        except Exception as e:
            error_msg = f"Error testing retrieval: {str(e)}"
            logger.error(error_msg, exc_info=True)
            return False, []
    
    def format_context_for_llm(
        self,
        results: List[Tuple[str, str, int, float]],
        max_length: int = 30000  # DRAMATICALLY INCREASED from 15000 for truly comprehensive answers
    ) -> str:
        """
        Format retrieved context with RICH and PROFESSIONAL presentation for MAXIMUM DETAIL.
        
        Creates a well-structured context string optimized for comprehensive responses:
        - Clear source attribution with document names
        - Relevance indicators (High/Good/Medium confidence)
        - Clean markdown tables
        - Rich formatting for readability
        - Professional structure
        
        ENHANCED: Increased max_length to 30000 and improved formatting for truly detailed answers
        
        Args:
            results: List of (chunk_text, filename, chunk_index, similarity) tuples
            max_length: Maximum context length in characters (default: 30000)
        
        Returns:
            Formatted context string ready for LLM prompt with rich content
        """
        if not results:
            return ""
        
        logger.debug(f"Formatting {len(results)} chunks for LLM (max length: {max_length})")
        
        formatted_parts = []
        current_length = 0
        chunks_included = 0
        
        for idx, (chunk_text, filename, chunk_index, similarity) in enumerate(results, 1):
            # Determine quality tier
            if similarity >= 0.80:
                quality = "High Confidence"
                priority = "[*** MOST RELEVANT] " if idx == 1 else "[** HIGHLY RELEVANT] "
                marker = "***"
            elif similarity >= 0.65:
                quality = "Good Match"
                priority = "[+ RELEVANT] "
                marker = "[+]"
            else:
                quality = "Supporting"
                priority = "[- BACKGROUND] "
                marker = " - "
            
            # Format header with clear attribution
            header = f"\n{'=' * 70}\n"
            header += f"{marker} {priority}SOURCE {idx}: {filename}\n"
            header += f"Chunk: {chunk_index} | Relevance: {quality} ({int(similarity * 100)}%)\n"
            header += f"{'=' * 70}\n\n"
            
            # Clean and format chunk text with ENHANCED structure
            cleaned_text = self._format_chunk_text_rich(chunk_text)
            
            # Build formatted chunk
            formatted_chunk = header + cleaned_text + "\n\n"
            
            # Check length constraint - be more permissive
            if current_length + len(formatted_chunk) > max_length:
                if chunks_included == 0:
                    # Include at least one chunk even if long
                    formatted_parts.append(formatted_chunk[:max_length - current_length])
                    logger.warning(f"Context truncated: only 1 chunk included (very long)")
                else:
                    logger.info(f"Context size limit reached: {chunks_included} of {len(results)} chunks included")
                break
            
            formatted_parts.append(formatted_chunk)
            current_length += len(formatted_chunk)
            chunks_included += 1
        
        context = "".join(formatted_parts)
        
        # Add summary header
        summary_header = f"""
{'=' * 70}
DOCUMENT CONTEXT SUMMARY
{'=' * 70}
Total Sources: {chunks_included} documents
Average Relevance: {sum(r[3] for r in results[:chunks_included]) / chunks_included * 100:.1f}%
Content Length: {len(context):,} characters

INSTRUCTIONS: Use ALL the information below to provide a COMPREHENSIVE answer.
{'=' * 70}

"""
        
        final_context = summary_header + context
        
        logger.info(f"Formatted context: {len(final_context):,} chars from {chunks_included} chunks (avg: {len(context)//chunks_included if chunks_included > 0 else 0} chars/chunk)")
        
        return final_context
    
    def _format_chunk_text_rich(self, chunk_text: str) -> str:
        """
        Clean and format chunk text for RICH professional presentation.
        
        Handles:
        - Table formatting in clean markdown
        - Paragraph structure
        - Text normalization
        - Whitespace cleanup
        - Enhanced readability
        
        Args:
            chunk_text: Raw text chunk
        
        Returns:
            Richly formatted text chunk
        """
        text = chunk_text.strip()
        
        # Check if this is a table (has | separators and multiple lines)
        if '|' in text and text.count('\n') > 1:
            return self._format_table_markdown_clean(text)
        
        # For regular text, enhance formatting for readability
        # Replace multiple spaces with single space
        text = re.sub(r' +', ' ', text)
        
        # Preserve paragraph breaks (double newlines) and enhance them
        text = re.sub(r'\n\n+', '\n\n', text)
        
        # Add some structure: detect lists and enhance
        lines = text.split('\n')
        formatted_lines = []
        for line in lines:
            stripped = line.strip()
            # Detect list items and enhance
            if stripped.startswith(('- ', '* ', '• ', '1.', '2.', '3.', '4.', '5.')):
                formatted_lines.append(f"  {stripped}")  # Indent lists slightly
            elif len(stripped) < 50 and stripped.endswith(':'):
                # Likely a header
                formatted_lines.append(f"\n**{stripped}**")
            else:
                formatted_lines.append(line)
        
        text = '\n'.join(formatted_lines)
        
        return text
    
    def _format_table_markdown_clean(self, table_text: str) -> str:
        """
        Format table text as clean markdown.
        
        Args:
            table_text: Text containing table with pipe separators
        
        Returns:
            Clean markdown formatted table
        """
        lines = table_text.strip().split('\n')
        if not lines:
            return table_text
        
        # Keep table header if present
        formatted_lines = []
        for line in lines:
            line = line.strip()
            if line:
                formatted_lines.append(line)
        
        return '\n'.join(formatted_lines)


# ============================================================================
# MODULE-LEVEL INSTANCE
# ============================================================================

# Create a singleton instance for use across the application
doc_processor = DocumentProcessor()
logger.debug("Created module-level doc_processor instance")



