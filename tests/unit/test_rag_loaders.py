
"""
Unit Tests for RAG Document Loaders
====================================

Comprehensive tests for document loading functionality in src/rag.py

Target: Increase coverage from 63% to 85% (+7% total coverage)

Focus areas:
- PDF loading (with pdfplumber and PyPDF2)
- DOCX loading
- Text file loading
- Error handling
- Content validation

Author: LocalChat Team
Created: January 2025
"""

from collections.abc import Callable
from pathlib import Path
from unittest.mock import patch

import pytest


@pytest.fixture
def real_pdf(tmp_path) -> Callable[[str], Path]:
    """Build an actual one-page PDF containing *text*."""
    def _build(text: str) -> Path:
        from reportlab.pdfgen import canvas

        path = tmp_path / "sample.pdf"
        pdf = canvas.Canvas(str(path))
        pdf.drawString(72, 720, text)
        pdf.save()
        return path
    return _build


@pytest.fixture
def blank_pdf(tmp_path) -> Path:
    """A structurally valid PDF with no text on it."""
    from reportlab.pdfgen import canvas

    path = tmp_path / "blank.pdf"
    pdf = canvas.Canvas(str(path))
    pdf.showPage()
    pdf.save()
    return path


@pytest.fixture
def real_docx(tmp_path) -> Path:
    """An actual .docx with two paragraphs and a table, since the loader reads both."""
    from docx import Document as DocxDocument

    path = tmp_path / "sample.docx"
    document = DocxDocument()
    document.add_paragraph("Paragraph 1 content")
    document.add_paragraph("Paragraph 2 content")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Region"
    table.cell(0, 1).text = "Benelux"
    document.save(str(path))
    return path

import pytest


class TestTextFileLoading:
    """Test plain text file loading."""

    def test_load_text_file_returns_content(self, tmp_path):
        """Test successful text file loading."""
        from src.rag import doc_processor

        # Create test file
        test_file = tmp_path / "test.txt"
        test_content = "This is test content.\\nMultiple lines.\\n"
        test_file.write_text(test_content, encoding='utf-8')

        success, content = doc_processor.load_text_file(str(test_file))

        assert success is True
        assert content == test_content
        assert len(content) > 0

    def test_load_text_file_handles_encoding(self, tmp_path):
        """Test text file with UTF-8 encoding."""
        from src.rag import doc_processor

        test_file = tmp_path / "unicode.txt"
        test_content = "Unicode: special characters test"
        test_file.write_text(test_content, encoding='utf-8')

        success, content = doc_processor.load_text_file(str(test_file))

        assert success is True
        assert "special" in content

    def test_load_text_file_handles_missing_file(self):
        """Test loading non-existent file."""
        from src.rag import doc_processor

        success, error = doc_processor.load_text_file("nonexistent.txt")

        assert success is False
        assert "No such file" in error or "not found" in error.lower()


class TestPDFLoading:
    """Test PDF file loading against real PDFs.

    These were skipped as "pdfplumber uses dynamic imports, difficult to mock",
    which diagnosed the wrong problem: the loader picks an extractor at runtime
    precisely so it can fall back, and a test that mocks the extractor asserts
    the mock. `reportlab` is already a dependency, so the fixture can be a real
    PDF and the assertion can be about what the loader actually extracted.
    """

    def test_load_pdf_returns_the_text_in_the_file(self, real_pdf):
        from src.rag import doc_processor

        success, content = doc_processor.load_pdf_file(str(real_pdf("Quarterly revenue summary")))

        assert success is True
        assert "Quarterly revenue summary" in content

    def test_load_pdf_reports_failure_for_a_pdf_with_no_text(self, blank_pdf):
        """The negative side: a PDF that parses fine but yields nothing."""
        from src.rag import doc_processor

        success, error = doc_processor.load_pdf_file(str(blank_pdf))

        assert success is False
        assert "empty" in error.lower() or "no text" in error.lower()


class TestDOCXLoading:
    """Test DOCX file loading."""

    def test_load_docx_returns_every_paragraph(self, real_docx):
        from src.rag import doc_processor

        success, content = doc_processor.load_docx_file(str(real_docx))

        assert success is True
        assert "Paragraph 1 content" in content
        assert "Paragraph 2 content" in content

    def test_load_docx_also_returns_table_cells(self, real_docx):
        """The docstring promises tables as well as paragraphs, and the mocked
        version set `tables = []` — so the half that reads tables was asserted
        by nothing."""
        from src.rag import doc_processor

        success, content = doc_processor.load_docx_file(str(real_docx))

        assert success is True
        assert "Benelux" in content

    def test_load_docx_handles_missing_file(self):
        """Test handling of missing DOCX file."""
        from src.rag import doc_processor

        with patch('src.rag.Document', side_effect=FileNotFoundError("File not found")):
            success, error = doc_processor.load_docx_file("missing.docx")

            assert success is False
            assert "not found" in error.lower()


class TestDocumentLoading:
    """Test main document loading function."""

    def test_load_document_routes_to_text(self, tmp_path):
        """Test load_document routes .txt files correctly."""
        from src.rag import doc_processor

        test_file = tmp_path / "test.txt"
        test_file.write_text("Text content", encoding='utf-8')

        success, content = doc_processor.load_document(str(test_file))

        assert success is True
        assert "Text content" in content

    def test_load_document_routes_to_pdf(self, real_pdf):
        """Routing, asserted on the routed-to loader's actual output rather than
        on `len(content) > 0`, which the text loader would satisfy too."""
        from src.rag import doc_processor

        success, content = doc_processor.load_document(str(real_pdf("Routed to the PDF loader")))

        assert success is True
        assert "Routed to the PDF loader" in content

    def test_load_document_handles_unsupported_type(self):
        """Test load_document rejects unsupported file types."""
        from src.rag import doc_processor

        success, error = doc_processor.load_document("test.xyz")

        assert success is False
        assert "unsupported" in error.lower() or "file type" in error.lower()


class TestImageFileLoading:
    """Test image file loading via vision model."""

    def test_load_image_file_fails_without_vision_model(self, tmp_path):
        """No vision model should return a clear error without reading the file."""
        from src.rag import doc_processor

        img_file = tmp_path / "photo.jpg"
        img_file.write_bytes(b"\xff\xd8\xff" + b"\x00" * 10)

        with patch('src.rag.loaders.ollama_client') as mock_client:
            mock_client.get_vision_model.return_value = None
            success, error = doc_processor.load_image_file(str(img_file))

        assert success is False
        assert "vision model" in error.lower()
        mock_client.get_vision_model.assert_called_once()

    def test_load_image_file_skips_file_read_when_no_vision_model(self, tmp_path):
        """Vision model check must happen before reading the file into memory."""
        from src.rag import doc_processor

        img_file = tmp_path / "large.png"
        img_file.write_bytes(b"\x89PNG" + b"\x00" * 100)

        with patch('src.rag.loaders.ollama_client') as mock_client:
            mock_client.get_vision_model.return_value = None
            with patch('builtins.open', side_effect=AssertionError("file was read")) as mock_open:
                success, error = doc_processor.load_image_file(str(img_file))
                # open() must not have been called for the image data read
                mock_open.assert_not_called()

        assert success is False

    def test_load_image_file_describes_image_with_vision_model(self, tmp_path):
        """When a vision model is present, the description is returned."""
        from src.rag import doc_processor

        img_file = tmp_path / "cat.jpg"
        img_file.write_bytes(b"\xff\xd8\xff" + b"\x00" * 10)

        with patch('src.rag.loaders.ollama_client') as mock_client:
            mock_client.get_vision_model.return_value = 'llava'
            mock_client.describe_image.return_value = (True, "A photo of a cat.")
            success, content = doc_processor.load_image_file(str(img_file))

        assert success is True
        assert "cat.jpg" in content
        assert "A photo of a cat." in content


# ---------------------------------------------------------------------------
# Excel loading
# ---------------------------------------------------------------------------

class TestExcelFileLoading:
    """Tests for load_excel_file via openpyxl."""

    def test_load_excel_returns_sheet_content(self, tmp_path):
        """Happy path: single sheet with data is returned as pipe-delimited text."""
        import openpyxl

        from src.rag import doc_processor

        path = tmp_path / "data.xlsx"
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Sales"
        ws.append(["Product", "Q1", "Q2"])
        ws.append(["Widget", 100, 200])
        wb.save(str(path))

        success, content = doc_processor.load_excel_file(str(path))

        assert success is True
        assert "Sales" in content
        assert "Product" in content
        assert "Widget" in content

    def test_load_excel_multiple_sheets(self, tmp_path):
        """Multiple sheets are all included in the output."""
        import openpyxl

        from src.rag import doc_processor

        path = tmp_path / "multi.xlsx"
        wb = openpyxl.Workbook()
        ws1 = wb.active
        ws1.title = "Sheet1"
        ws1.append(["A", "B"])
        ws2 = wb.create_sheet("Sheet2")
        ws2.append(["X", "Y"])
        wb.save(str(path))

        success, content = doc_processor.load_excel_file(str(path))

        assert success is True
        assert "Sheet1" in content
        assert "Sheet2" in content

    def test_load_excel_empty_workbook_fails(self, tmp_path):
        """A workbook with no data returns failure."""
        import openpyxl

        from src.rag import doc_processor

        path = tmp_path / "empty.xlsx"
        wb = openpyxl.Workbook()
        wb.active.title = "Empty"
        wb.save(str(path))

        success, content = doc_processor.load_excel_file(str(path))

        assert success is False
        assert "no extractable data" in content.lower()

    def test_load_excel_unavailable_returns_error(self, tmp_path):
        """When openpyxl is not installed the loader returns a clear error."""
        from unittest.mock import patch

        from src.rag import doc_processor

        path = tmp_path / "x.xlsx"
        path.write_bytes(b"PK\x03\x04")  # not a real xlsx, but loader won't reach it

        with patch("src.rag.loaders.XLSX_AVAILABLE", False):
            success, error = doc_processor.load_excel_file(str(path))

        assert success is False
        assert "openpyxl" in error.lower()
